
from docx import Document
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class CoverLetter:
    """
    Fields:
        - first_name (str)
        - last_name (str)
        - address (str): including building/APT number
        - city (str)
        - state (str)
        - zip (str)
        - email (str)
        - phone_number (str): any format
        - company_name (str)
        - job_title (str)
        - unique_para (str): User-defined unique paragraph extension for the specific role. 
            (f"I am excited to apply for the {self.job_title} position at {self.company_name} 
            is already the first line of this paragraph. 
        - body (list[str]): strings for each paragraph 
        
    
    Returns: An automated customized cover letter in Microsoft Word using Python.
    """

    def __init__(self, first_name, last_name, address, city, state, zip, email, phone_number, 
                 company_name, job_title, body, unique_para):
        
        self.first_name = first_name
        self.last_name = last_name
        self.address = address
        self.city = city
        self.state = state
        self.zip = zip
        self.email = email
        self.phone_number = phone_number
        self.company_name = company_name
        self.job_title = job_title
        self.body = body
        self.unique_para = unique_para
        

    def generate_cover_letter(self):

        doc = Document() # initializing the word document 
        heading = doc.add_heading(f'{self.first_name} {self.last_name}', level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center heading
        contact_info = doc.add_paragraph(
            f'{self.address} {self.city}, {self.state} {self.zip} | {self.phone_number} | {self.email}'
        )
        contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
        current_date = datetime.today().strftime("%B %d, %Y")  # e.g., "July 5, 2025"
        doc.add_paragraph(f'\n\n{current_date}') 
        doc.add_paragraph(f'Dear Hiring Manager at {self.company_name},')
        doc.add_paragraph(f"\nI am excited to apply for the {self.job_title} position at {self.company_name}. {self.unique_para}")
        for p in self.body: 
            doc.add_paragraph(p)
        doc.add_paragraph(f"I look forward to bringing my analytical skills and collaborative approach to {self.company_name}. I welcome the opportunity to discuss how my experience can contribute to your team. Thank you for your time and consideration.")
        doc.add_paragraph('\n\nSincerely,')
        doc.add_paragraph(f'\n{self.first_name} {self.last_name}')
        # 
        doc.save(f'{self.first_name}{self.last_name}{self.company_name}CoverLetter.docx')



          
    

# child class from CoverLetter
# make sure you run CoverLetter first so that it's in memory 
class CustomCoverLetter(CoverLetter):
    """custom cover letters for specific individuals for different jobs."""

    def __init__(self, first_name, last_name, address, city, state, zip, email, phone_number):
        super().__init__(first_name, last_name, address, city, state, zip, email, phone_number,
                         company_name=None, job_title=None, body=[], unique_para="")

    def set_job_details(self, company_name, job_title, unique_para, body):
        """Stores your personal data in a child class, so that you only need to input new information."""
        self.company_name = company_name
        self.job_title = job_title
        self.unique_para = unique_para
        self.body = body



# generic example below

## new object from the CustomCoverLetter
# salty_grad_student = CustomCoverLetter(
#     first_name="Salty",
#     last_name="GradStudent",
#     address="123 Main St, Apt 4B",
#     city="Kansas City",
#     state="Ks",
#     zip="66103",
#     email="salty@kumc.edu",
#     phone_number="123-456-7890"
# ) # if using the same body over and over again then you can store it here too 

# # Let's you switch up your body paragraphs 
# # Copying an pasting from an actual generic cover letter that you wrote for this kind
# # of role is recommended 
# saltys_ds_body = [ # obviously put more info than this... each element is a paragraph of the cover letter 
#         "I have experience working with scalable software systems.",
#         "I am proficient in Python, R, and SQL.",
#         "I am passionate about data-driven solutions and problem-solving."
#     ]
# saltys_bioinformatics_body = [
#     "I don't know biology.", 
#     "What is a protien?", 
#     "Why did I major in history? "
# ]



# salty_grad_student.set_job_details(
#     company_name="Tech Corp",
#     job_title="Data Scientist",
#     unique_para="I have a strong background in Python, R, statistics and automation.",
#     body=saltys_ds_body
# )
# salty_grad_student.generate_cover_letter() # makes an example word doc of a cover letter
